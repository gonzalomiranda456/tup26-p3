from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path("CURRICULUM VITAE - Alejandro Di Battista.docx")
ACCENT = RGBColor(196, 85, 18)
GRAY = RGBColor(95, 95, 95)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def clear_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def p(cell_or_doc, text="", size=10, bold=False, color=None, align=None, after=3):
    para = cell_or_doc.add_paragraph()
    para.paragraph_format.space_after = Pt(after)
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if align:
        para.alignment = align
    return para


def heading(cell_or_doc, text):
    para = p(cell_or_doc, text.upper(), size=14, bold=True, color=GRAY, after=7)
    para.paragraph_format.space_before = Pt(8)
    return para


def bullet(cell_or_doc, text, size=10):
    para = cell_or_doc.add_paragraph(style=None)
    para.style = "List Bullet"
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    return para


def job(cell_or_doc, title, org, dates, bullets):
    p(cell_or_doc, title.upper(), size=11, bold=True, color=RGBColor(40, 40, 40), after=1)
    p(cell_or_doc, org, size=10.5, bold=True, after=1)
    p(cell_or_doc, dates, size=10, color=GRAY, after=2)
    for item in bullets:
        bullet(cell_or_doc, item, size=9.5)
    p(cell_or_doc, "", after=2)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.1)
section.bottom_margin = Cm(1.1)
section.left_margin = Cm(1.2)
section.right_margin = Cm(1.2)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10)

layout = doc.add_table(rows=1, cols=2)
layout.autofit = False
layout.columns[0].width = Cm(6.4)
layout.columns[1].width = Cm(11.8)
clear_borders(layout)
left = layout.cell(0, 0)
right = layout.cell(0, 1)
set_cell_shading(left, "D9D9D9")

for cell in (left, right):
    cell.margin_top = Cm(0.2)
    cell.margin_bottom = Cm(0.2)
    cell.margin_left = Cm(0.35)
    cell.margin_right = Cm(0.35)

p(right, "ALEJANDRO DI", size=28, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
p(right, "BATTISTA", size=28, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
p(right, "Asesor Tecnologico", size=22, bold=True, color=RGBColor(130, 130, 130), align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

heading(left, "Datos personales")
p(left, "DNI: 18.627.585")
p(left, "Fecha de nacimiento: 5 de diciembre de 1967")
p(left, "Edad: 58 años")
p(left, "Nacionalidad: argentina")
p(left, "Estado civil: divorciado", after=12)

heading(left, "Contacto")
p(left, "alejandrodibattista@gmail.com")
p(left, "381 5343458")
p(left, "Argentina")
p(left, "www.linkedin.com/in/alejandrodibattista", after=12)

heading(left, "Perfil")
p(
    left,
    "Integrando la pasión por emprender y la tecnología. Enfocado completamente en el desarrollo "
    "de productos mobile, con experiencia en asesoramiento tecnológico, desarrollo de software y docencia universitaria.",
    after=12,
)

heading(left, "Habilidades")
for skill in ["C#", "Flutter", "Dart", "Python", "JavaScript", "Desarrollo mobile"]:
    bullet(left, skill)

heading(left, "Idiomas")
bullet(left, "English - Middle (1100-1500)")
bullet(left, "Español - Nativo")

heading(right, "Experiencia laboral")
job(
    right,
    "Gerente de Sistema",
    "Caja Popular de Ahorros de Tucumán",
    "Septiembre 2021 - Septiembre 2022",
    ["Gestión de sistemas en San Miguel de Tucumán, Tucumán, Argentina."],
)
job(
    right,
    "Asesor Tecnologico",
    "Concejo Deliberante Yerba Buena",
    "Marzo 2017 - Diciembre 2023",
    ["Asesoramiento tecnológico en Yerba Buena, Tucumán, Argentina."],
)
job(
    right,
    "Director Proyecto Modernización de Hospitales de Tucumán",
    "UTN",
    "Enero 2003 - Diciembre 2006",
    [
        "Dirección integral del proyecto.",
        "Elaboración de pliegos de licitación.",
        "Supervisión de construcciones.",
        "Formación de equipos de desarrollo y gestión.",
        "Ejecución del proyecto.",
    ],
)

heading(right, "Experiencia docente")
job(
    right,
    "Profesor universitario",
    "Universidad Tecnológica Nacional (UTN)",
    "Marzo 2024 - Actualidad",
    ["Enseñanza de programación en C#, Python y JavaScript."],
)
job(
    right,
    "Profesor Introducción a la Comercialización",
    "Universidad Santo Tomás de Aquino",
    "1998 - 2010",
    ["Profesor adjunto."],
)

doc.add_section(WD_SECTION.NEW_PAGE)
section = doc.sections[-1]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.1)
section.bottom_margin = Cm(1.1)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)

p(doc, "ALEJANDRO DI BATTISTA", size=18, bold=True, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

heading(doc, "Experiencia docente")
job(
    doc,
    "Profesor Análisis de Desempeño de Sistema de Procesamiento de Datos",
    "Universidad Católica de Santiago del Estero",
    "1999 - 2006",
    ["Profesor titular."],
)
job(
    doc,
    "Profesor de Informática",
    "Universidad Santo Tomás de Aquino",
    "1997 - 2005",
    ["Profesor titular."],
)

heading(doc, "Otras experiencias")
job(
    doc,
    "Desarrollo de Personal",
    "Algacom Argentina",
    "Febrero 2018 - Marzo 2019",
    ["Experiencia en desarrollo de personal."],
)
job(
    doc,
    "Propietario",
    "AgilPad",
    "Septiembre 2011 - Diciembre 2016",
    ["CEO en San Miguel de Tucumán, Tucumán, Argentina."],
)
job(
    doc,
    "Propietario",
    "Agilsoft",
    "Marzo 2006 - Septiembre 2011",
    ["Empresa dedicada al desarrollo a medida de aplicaciones web."],
)
job(
    doc,
    "Asesor Tecnologico",
    "IPLA",
    "Enero 2009 - Mayo 2009",
    ["Implementación de la gestión online del servicio."],
)
job(
    doc,
    "Capacitación",
    "Scania Latin America",
    "1994 - 1996",
    ["Capacitación en herramientas de gestión a la gerencia."],
)

heading(doc, "Educación")
job(
    doc,
    "Ingeniería en Sistemas de Información",
    "Universidad Tecnológica Nacional",
    "1988 - 1994",
    [],
)
job(
    doc,
    "MBA",
    "Fundación del Tucumán",
    "1995 - 1997",
    [],
)

doc.save(OUT)
print(OUT.resolve())
